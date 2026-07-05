#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
make_word.py
توليد نسخ Word (.docx) قابلة للتعديل وسهلة القراءة (عربية RTL) من المخرجات:
  1) تحويل تقارير Markdown إلى Word (عناوين/قوائم/جداول/اقتباسات).
  2) مستند Word موحّد بالجداول الرئيسية (سجل الصكوك، المحطات، الأسباب، المبالغ، الجرد).
  3) مستند Word موحّد بكل النصوص المستخرجة من المستندات، مرتّبة حسب الأقسام.
المخرجات في: outputs/word/
الاستخدام: python3 make_word.py --staging <staging_dir>
"""
import argparse
import csv
import json
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# فُكّ الارتباط بوحدة القضية (audit_case): الثوابت الوحيدة المطلوبة تُعرَّف هنا
# بنفس الدلالات — الجذر من CASE_ROOT أو مجلد الحزمة، بلا أي منطق قضية.
import os as _os
OUTPUT_ROOT = Path(_os.environ.get("CASE_ROOT") or Path(__file__).resolve().parent.parent)
DEFAULT_STAGING = OUTPUT_ROOT / "staging"

OUT_MD = OUTPUT_ROOT / "outputs" / "markdown"
OUT_CSV = OUTPUT_ROOT / "outputs" / "csv"
OUT_JSON = OUTPUT_ROOT / "outputs" / "json"
OUT_WORD = OUTPUT_ROOT / "outputs" / "word"
OUT_WORD.mkdir(parents=True, exist_ok=True)

TAG_S = "[مخرج آلي / Script Output - سماوي]"
TAG_R = "[مخرج آلي يحتاج مراجعة بشرية - وردي]"


def set_rtl(paragraph):
    """يجعل الفقرة من اليمين لليسار (RTL) ومحاذاة لليمين."""
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement('w:bidi')
    pPr.append(bidi)


def style_doc(doc):
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(12)
    # ضبط الخط العربي
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts()
    rfonts.set(qn('w:cs'), 'Arial')


_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def xml_safe(s):
    """يزيل المحارف التحكّمية غير المتوافقة مع XML (تظهر من OCR) لتفادي أخطاء Word/Excel."""
    return _CTRL.sub("", s) if isinstance(s, str) else s


def add_heading_rtl(doc, text, level=1):
    h = doc.add_heading(xml_safe(text), level=level)
    set_rtl(h)
    return h


def add_para_rtl(doc, text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    set_rtl(p)
    run = p.add_run(xml_safe(text))
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    return p


def md_to_docx(md_path, docx_path, title=None):
    if not md_path.exists():
        return False
    doc = Document()
    style_doc(doc)
    lines = md_path.read_text(encoding="utf-8").split("\n")
    table_buf = []

    def flush_table():
        nonlocal table_buf
        if not table_buf:
            return
        # صفوف الجدول، تجاهل سطر الفاصل |---|
        rows = [r for r in table_buf if not re.match(r"^\s*\|?[\s:\-\|]+\|?\s*$", r)]
        cells = [[c.strip() for c in r.strip().strip("|").split("|")] for r in rows]
        if cells:
            ncol = max(len(r) for r in cells)
            t = doc.add_table(rows=0, cols=ncol)
            t.style = "Light Grid Accent 1"
            t.alignment = 2  # right
            for r in cells:
                cellsrow = t.add_row().cells
                for i in range(ncol):
                    txt = r[i] if i < len(r) else ""
                    cellsrow[i].text = xml_safe(txt)
                    for para in cellsrow[i].paragraphs:
                        set_rtl(para)
        table_buf = []

    for ln in lines:
        s = ln.rstrip()
        if s.strip().startswith("|"):
            table_buf.append(s)
            continue
        else:
            flush_table()
        if not s.strip():
            continue
        if s.startswith("#"):
            m = re.match(r"^(#+)\s*(.*)", s)
            lvl = min(len(m.group(1)), 4)
            add_heading_rtl(doc, m.group(2), level=lvl)
        elif s.startswith(">"):
            add_para_rtl(doc, s.lstrip("> ").strip(), italic=True,
                         color=RGBColor(0x80, 0x00, 0x00))
        elif re.match(r"^\s*[-*]\s+", s):
            txt = re.sub(r"^\s*[-*]\s+", "", s)
            txt = txt.replace("**", "")
            p = doc.add_paragraph(txt, style="List Bullet")
            set_rtl(p)
        elif s.strip() == "---":
            doc.add_paragraph("")
        else:
            add_para_rtl(doc, s.replace("**", ""))
    flush_table()
    doc.save(str(docx_path))
    return True


def csv_to_table(doc, csv_path, title, max_rows=400):
    if not csv_path.exists():
        return
    add_heading_rtl(doc, title, level=2)
    with open(csv_path, encoding="utf-8-sig") as f:
        rows = list(csv.reader(f))
    if not rows:
        return
    header = rows[0]
    # احذف عمودي الوسم لتبسيط القراءة (نضيف ملاحظة عامة بدلاً منها)
    keep = [i for i, h in enumerate(header) if "وسم" not in h and "مفتاح" not in h]
    t = doc.add_table(rows=1, cols=len(keep))
    t.style = "Light Grid Accent 1"
    t.alignment = 2
    for ci, i in enumerate(keep):
        t.rows[0].cells[ci].text = header[i]
        for para in t.rows[0].cells[ci].paragraphs:
            set_rtl(para)
    for r in rows[1:max_rows + 1]:
        cells = t.add_row().cells
        for ci, i in enumerate(keep):
            cells[ci].text = xml_safe((r[i] if i < len(r) else "")[:300])
            for para in cells[ci].paragraphs:
                set_rtl(para)
    doc.add_paragraph("")


def build_tables_doc():
    doc = Document()
    style_doc(doc)
    add_heading_rtl(doc, "الجداول الرئيسية للقضية", level=0)
    add_para_rtl(doc, f"{TAG_S}  {TAG_R}", italic=True)
    add_para_rtl(doc, "جداول آلية تحتاج مراجعة بشرية — لا تُعد رأياً قانونياً نهائياً.", italic=True)
    for csvname, title in [
        ("12_deeds_register.csv", "سجلّ الصكوك/الأحكام"),
        ("13_key_milestones.csv", "المحطات الرئيسية (مرتّبة)"),
        ("09_central_amounts.csv", "المبالغ المركزية"),
        ("10_objection_cassation_grounds.csv", "أسباب الاعتراض/النقض والدفوع"),
        ("11_central_law_references.csv", "الإحالات النظامية المركزية"),
        ("01_full_file_inventory.csv", "الجرد الكامل للملفات"),
        ("06_unreadable_or_ocr_needed.csv", "الملفات غير المقروءة/تحتاج OCR"),
    ]:
        csv_to_table(doc, OUT_CSV / csvname, title)
    doc.save(str(OUT_WORD / "10_all_key_tables.docx"))


def build_full_text_doc(staging):
    """مستند Word موحّد بكل النصوص المستخرجة، مرتّبة حسب الأقسام."""
    data = json.loads((OUT_JSON / "full_audit_data.json").read_text(encoding="utf-8"))
    files = data["files"]
    tdir = Path(staging) / "text"
    rdir = Path(staging) / "text_readable"  # النص المُصحَّح الترتيب إن توفّر
    # رتّب حسب القسم ثم ترتيب الدراسة
    files.sort(key=lambda f: (f.get("parent_path", ""), f.get("title", "")))
    doc = Document()
    style_doc(doc)
    add_heading_rtl(doc, "النصوص المستخرجة من مستندات القضية", level=0)
    add_para_rtl(doc, f"{TAG_S}  {TAG_R}", italic=True)
    add_para_rtl(doc, "تجميع آلي لنصوص المستندات المقروءة. النصوص الممسوحة ضوئياً (OCR) "
                      "قد يكون ترتيبها بصرياً ويصعب قراءتها؛ موسومة بذلك. تحتاج مراجعة.",
                 italic=True)
    fixed_ids = set()
    fx = rdir / "_fixed_ids.json"
    if fx.exists():
        fixed_ids = set(json.loads(fx.read_text(encoding="utf-8")))
    cur_section = None
    n = 0
    for f in files:
        rp = rdir / f"{f['id']}.txt"
        p = rp if rp.exists() else (tdir / f"{f['id']}.txt")
        if not p.exists():
            continue
        sec = f.get("parent_path", "") or "—"
        if sec != cur_section:
            cur_section = sec
            add_heading_rtl(doc, f"القسم: {sec}", level=1)
        add_heading_rtl(doc, f.get("title", "")[:140], level=2)
        meta = f"النوع: {f.get('doc_type','')} | اتجاه النص: {f.get('orientation','')} | الرابط: {f.get('viewUrl','')}"
        add_para_rtl(doc, meta, italic=True, color=RGBColor(0x55, 0x55, 0x55))
        txt = p.read_text(encoding="utf-8", errors="replace")
        if f["id"] in fixed_ids:
            add_para_rtl(doc, "[ملاحظة: نص ممسوح ضوئياً (OCR) صُحِّح ترتيبه آلياً ليصبح مقروءاً — "
                              "قد يحتوي أخطاء OCR طفيفة ويحتاج تدقيقاً]",
                         italic=True, color=RGBColor(0xC0, 0x00, 0x00))
        # قسّم النص لفقرات معقولة
        for chunk in re.split(r"\n{2,}", txt):
            chunk = chunk.strip()
            if chunk:
                add_para_rtl(doc, chunk[:4000])
        doc.add_page_break()
        n += 1
    doc.save(str(OUT_WORD / "11_extracted_documents_text.docx"))
    return n


def main():
    ap = argparse.ArgumentParser()
    ap.add_argument("--staging", default=str(DEFAULT_STAGING))
    args = ap.parse_args()

    # 1) تقارير Markdown -> Word
    reports = [
        ("00_case_brief.md", "00_ملخص_القضية.docx"),
        ("initial_audit_report.md", "01_تقرير_الفحص_الأولي.docx"),
        ("case_timeline_and_grounds.md", "02_الخط_الزمني_وأسباب_النقض.docx"),
        ("human_review_checklist.md", "03_قائمة_مراجعة_المحامي.docx"),
    ]
    done = 0
    for src, dst in reports:
        if md_to_docx(OUT_MD / src, OUT_WORD / dst):
            done += 1

    # 2) جداول Word
    build_tables_doc()

    # 3) نصوص المستندات Word
    n = build_full_text_doc(args.staging)

    print(f"حُوّلت {done} تقارير إلى Word + مستند الجداول + مستند النصوص ({n} مستنداً).")
    print(f"المخرجات في: {OUT_WORD}")
    for p in sorted(OUT_WORD.glob('*.docx')):
        print(" -", p.name)


if __name__ == "__main__":
    main()
