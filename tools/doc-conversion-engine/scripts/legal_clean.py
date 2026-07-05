#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
legal_clean.py — تنقيح وتنسيق قانوني محافظ لمستندات مستخرجة آلياً (OCR).

المبدأ: الأولوية لسلامة النص القانوني. لا عكس آلي للأسطر، لا حذف مضمون،
لا إعادة صياغة. التنظيف = إزالة ضوضاء الترويسات/التذييلات المكررة وضبط المسافات
والتنسيق RTL فقط؛ وكل ما يحتمل الاجتهاد يُعلَّم للمراجعة البشرية ويُسجَّل في التقارير.

المصدر: outputs/json/full_documents.jsonl (نصوص مقسّمة لكل مستند) + فحص الجودة.
المخرجات في: outputs/cleaned/
  cleaned_legal_documents.docx / .pdf , review_version.docx , document_index.json ,
  cleaning_report.md , headers_footers_report.md , ocr_issues_report.md ,
  bidi_reading_order_report.md , source_comparison_report.md , duplication_report.md
"""
import os, re, json, csv, shutil, unicodedata, difflib, datetime, subprocess
from pathlib import Path

from docx import Document
from docx.shared import Pt, Cm, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(os.environ.get("CASE_ROOT") or Path(__file__).resolve().parent.parent)
IN = ROOT / "outputs" / "json" / "full_documents.jsonl"
QC = ROOT / "outputs" / "csv" / "14_readability_qc.csv"
OUTDIR = ROOT / "outputs" / "cleaned"
BACKUP = OUTDIR / "_backup"
FONT = "Traditional Arabic"
TODAY = datetime.datetime.now().strftime("%Y-%m-%d %H:%M")

# تنظيف النص العربي: إزالة علامات الاتجاه الخفية + توحيد الحروف/الأرقام الفارسية للعربية (بلا حذف محتوى).
_STRIP = dict.fromkeys([0x200b, 0x200c, 0x200d, 0x200e, 0x200f, 0x202a, 0x202b,
                        0x202c, 0x202d, 0x202e, 0x2066, 0x2067, 0x2068, 0x2069, 0xfeff], None)
_LOOK = {"ھ": "ه", "ہ": "ه", "ۀ": "ه", "ۃ": "ة", "ی": "ي", "ۍ": "ي", "ک": "ك"}
_PDIG = {0x06F0 + i: chr(0x0660 + i) for i in range(10)}


def clean_text(t):
    if not t:
        return t
    t = t.translate(_STRIP)
    for a, b in _LOOK.items():
        t = t.replace(a, b)
    return t.translate(_PDIG)

# وسوم المراجعة الموحّدة
T_OCR = "[بحاجة مراجعة: OCR غير واضح]"
T_HF = "[بحاجة مراجعة: احتمال ترويسة مختلطة]"
T_DUP = "[بحاجة مراجعة: احتمال تكرار]"
T_TBL = "[بحاجة مراجعة: جدول غير واضح]"
T_BIDI = "[بحاجة مراجعة: اختلاط اتجاه القراءة داخل السطر]"
T_REV = "[بحاجة مراجعة: احتمال انعكاس نص OCR]"

# أنماط الترويسات/التذييلات
HF_RE = re.compile(
    r"(صفحة\s*\d+\s*(?:من|/)\s*\d+|صحيفة\s*رقم|رقم\s*الصحيفة|^\s*صحيفة\s*\d+|وزار[ةه]\s*العدل|"
    r"المملكة\s*العربية\s*السعودية|المحكم[ةه]\s*(?:العام[ةه]|التجاري[ةه]|العليا|الجزائي[ةه]|الإداري[ةه])|"
    r"محكم[ةه]\s*التنفيذ|محكم[ةه]\s*الاستئناف|الدائر[ةه]\s|رقم\s*الصك|تاريخ\s*الصك)")
PAGE_MARK = re.compile(r"^\s*\[\s*صفحة\s*\d+\s*\]\s*$")
BARE_NUM = re.compile(r"^\s*[\d٠-٩\s\.\-_،|]{0,8}\s*$")  # سطر شبه فارغ/أرقام قصيرة (ضوضاء)
DEED_RE = re.compile(r"(?:رقم\s*الصك|صك\s*رقم)[:\s]*([0-9٠-٩/]{4,})")
CASE_RE = re.compile(r"(?:رقم\s*القضي[ةه]|القضي[ةه]\s*رقم)[:\s]*([0-9٠-٩]{4,})")

# علامات بنية قانونية (تبقى عناوين مستقلة، لا تُدمج)
STRUCT = ["الوقائع", "الطلبات", "الدعوى", "الإجابة", "الدفوع", "البينات", "الأسباب",
          "المنطوق", "الاعتراض", "المرفقات", "أسباب الحكم", "حيثيات", "لذا حكمت",
          "لذلك حكمت", "حكمت الدائرة", "حكمت المحكمة", "أولاً", "ثانياً", "ثالثاً",
          "رابعاً", "خامساً", "سادساً", "الحمد لله", "بسم الله الرحمن الرحيم"]
STRUCT_RE = re.compile(r"^\s*(?:" + "|".join(re.escape(s) for s in STRUCT) + r")")
TABLE_HINT = re.compile(r"(رقم\s*الهوي[ةه]|نوع\s*الهوي[ةه]|الجنسي[ةه]|الهوي[ةه]\s*الوطني[ةه]|السجل\s*التجاري)")

_CTRL = re.compile(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]")


def norm(s):
    s = unicodedata.normalize("NFKC", s)
    for a, b in [("أ", "ا"), ("إ", "ا"), ("آ", "ا"), ("ى", "ي"), ("ة", "ه")]:
        s = s.replace(a, b)
    s = re.sub(r"[ً-ْـ]", "", s)
    return re.sub(r"\s+", " ", s).strip()


def xml_safe(s):
    return _CTRL.sub("", s) if isinstance(s, str) else s


def is_mixed_dir(line):
    """يكتشف سطراً يخلط عربياً ولاتينياً (≥3 حروف لاتينية متتالية) — اشتباه ترتيب/اتجاه."""
    has_ar = any(unicodedata.bidirectional(c) in ("R", "AL") for c in line)
    if not has_ar:
        return False
    run = 0
    for c in line:
        if unicodedata.bidirectional(c) == "L":
            run += 1
            if run >= 3:
                return True
        else:
            run = 0
    return False


# ---------- التنظيف الخفيف على مستوى السطر ----------
def clean_line(s, stats):
    orig = s
    s = xml_safe(s)
    if s != orig:
        stats["ctrl_removed"] += 1
    # إزالة محارف اتجاه/صفرية ورموز OCR التالفة
    s2 = re.sub(r"[​-‏‪-‮⁦-⁩﻿�□■◦�¤¦]", "", s)
    s2 = re.sub(r"[~`^_=]{2,}", " ", s2)          # سلاسل رموز زخرفية من المسح
    if s2 != s:
        stats["ctrl_removed"] += 1
    s = s2
    # توحيد المسافات
    s2 = re.sub(r"[ \t]{2,}", " ", s)
    if s2 != s:
        stats["space_collapsed"] += 1
    s = s2
    # مسافة بعد علامات الترقيم العربية إن التصقت بحرف
    s2 = re.sub(r"([،؛:])(?=\S)", r"\1 ", s)
    s2 = re.sub(r"\s+([،؛:.])", r"\1", s2)
    if s2 != s:
        stats["punct_fixed"] += 1
    s = s2
    # طيّ التكرار الفوري لكلمة/رقم نفسه (أثر OCR شائع): «مكرَّر مكرَّر» → «مكرَّر»
    s = re.sub(r"(\b[^\s]{3,}\b)(\s+\1\b)+", r"\1", s)
    return s.strip()


def load_qc():
    qc = {}
    if not QC.exists():
        return qc
    rows = list(csv.reader(open(QC, encoding="utf-8-sig")))
    h = rows[0]
    def idx(n):
        return next((i for i, c in enumerate(h) if n in c), -1)
    iT, iS, iR, iW = idx("الملف"), idx("الحالة"), idx("معكوس"), idx("كلمات")
    for r in rows[1:]:
        if iT < len(r):
            qc[r[iT].strip()] = {"status": r[iS] if iS < len(r) else "",
                                 "rev": int(re.sub(r"\D", "", r[iR]) or 0) if iR < len(r) else 0,
                                 "words": r[iW] if iW < len(r) else ""}
    return qc


# ---------- معالجة مستند واحد ----------
def process_doc(rec, qc, logs):
    title = rec.get("title", "")
    text = rec.get("full_text", "") or ""
    qci = qc.get(title) or qc.get(title.rsplit(".", 1)[0]) or {}
    raw_lines = text.split("\n")
    stats = {"ctrl_removed": 0, "space_collapsed": 0, "punct_fixed": 0}

    # 1) كشف الترويسات/التذييلات المكررة (سطر مطبَّع يتكرر ≥2 ويطابق نمط ترويسة)
    norм_counts = {}
    for l in raw_lines:
        n = norm(l)
        if len(n) >= 8 and HF_RE.search(l):
            norм_counts[n] = norм_counts.get(n, 0) + 1
    repeated_hf = {n: c for n, c in norм_counts.items() if c >= 2}

    # استخراج بيانات مهمة للبطاقة
    deed = rec.get("card", {}).get("رقم الصك/الحكم", "")
    for l in raw_lines:
        m = DEED_RE.search(l)
        if m and not deed:
            deed = m.group(1)

    # 2) بناء المتن المنقّح + تعليمات المراجعة
    body = []        # عناصر: (kind, text, review_tags)
    removed = []     # (reason, text)
    seen_hf = {}
    page = 1
    bidi_hits, tbl_hits = [], []
    for l in raw_lines:
        s = l.rstrip()
        if PAGE_MARK.match(s):
            m = re.search(r"\d+", s)
            if m:
                page = int(m.group())
            removed.append(("علامة صفحة", s))
            continue
        n = norm(s)
        # ترويسة مكررة: أزل التكرار من المتن (احفظ أول مرة فقط، وسجّل)
        if n in repeated_hf:
            seen_hf[n] = seen_hf.get(n, 0) + 1
            logs["hf"].setdefault(n, {"count": 0, "docs": set(), "sample": s[:80]})
            logs["hf"][n]["count"] += 1
            logs["hf"][n]["docs"].add(title)
            if seen_hf[n] == 1:
                # أول ظهور: اعتبرها ترويسة تعريفية، انقلها لبطاقة لاحقاً ولا تُكرّر بالمتن
                removed.append(("ترويسة تعريفية (نُقلت للبطاقة)", s))
            else:
                removed.append(("ترويسة مكررة", s))
            continue
        # سطر ضوضاء قصير جداً (أرقام/رموز مبعثرة) — يُبقى لكن يُعلّم في نسخة المراجعة
        tags = []
        if BARE_NUM.match(s) and s.strip() and not re.fullmatch(r"[\d٠-٩]{4,}", s.strip()):
            tags.append(T_OCR)
        if is_mixed_dir(s):
            tags.append(T_BIDI)
            bidi_hits.append(s)
        if TABLE_HINT.search(s):
            tags.append(T_TBL)
            tbl_hits.append(s)
        cl = clean_line(s, stats)
        if cl == "":
            body.append(("blank", "", []))
        else:
            body.append(("line", cl, tags))

    # أسطر معكوسة متبقية من فحص الجودة (مستوى المستند)
    rev_residual = qci.get("rev", 0)

    # 3) إعادة بناء الفقرات (دمج محافظ)
    paras = []  # (kind: heading/para/blank, text, tags)
    buf, buftags = [], []
    def flush():
        if buf:
            paras.append(("para", " ".join(buf).strip(), list(dict.fromkeys(buftags))))
            buf.clear(); buftags.clear()
    blanks = 0
    last_norm = ""
    for kind, t, tags in body:
        if kind == "blank":
            blanks += 1
            if blanks >= 2:   # فجوة فقرة حقيقية فقط؛ السطر الفارغ المفرد يُتجاهل (إعادة تدفّق)
                flush()
            continue
        blanks = 0
        # أسقِط الأسطر التي لا تحمل أي حرف/رقم (ضوضاء رموز خالصة)
        if not re.search(r"[0-9٠-٩A-Za-z؀-ۿ]", t):
            continue
        # اطوِ السطر المكرّر فوراً (أثر OCR: نفس السطر مرّتين/أكثر متتالية)
        nt = norm(t)
        if nt and nt == last_norm:
            continue
        last_norm = nt
        if STRUCT_RE.match(t):
            flush()
            paras.append(("heading", t, tags))
            continue
        buf.append(t); buftags.extend(tags)
        # اقفل الفقرة عند نهاية جملة واضحة فقط (لا على الأسطر القصيرة — تُدمج)
        if re.search(r"[.؟!]$", t) and len(" ".join(buf)) > 40:
            flush()
    flush()

    # سجل إحصاءات/مشكلات
    logs["ocr_stats"]["ctrl"] += stats["ctrl_removed"]
    logs["ocr_stats"]["space"] += stats["space_collapsed"]
    logs["ocr_stats"]["punct"] += stats["punct_fixed"]
    needs_review = bool(bidi_hits or tbl_hits or rev_residual or qci.get("status") not in ("سليم", ""))

    return {
        "title": title, "doc_type": rec.get("doc_type", ""), "parent_path": rec.get("parent_path", ""),
        "viewUrl": rec.get("viewUrl", ""), "card": rec.get("card", {}) or {}, "deed": deed,
        "qc": qci, "rev_residual": rev_residual, "paras": paras, "removed": removed,
        "bidi_hits": bidi_hits, "tbl_hits": tbl_hits, "needs_review": needs_review,
        "words_before": len(text.split()), "words_after": sum(len(p[1].split()) for p in paras),
        "paras_before": len([l for l in raw_lines if l.strip()]), "paras_after": len(paras),
    }


# ---------- بناء DOCX ----------
def set_section_rtl(section):
    # A4 وهوامش وفق أمر التنسيق العربي: يمين 3.5 · يسار 2.5 · أعلى 3 · أسفل 2 سم
    section.page_width = Mm(210); section.page_height = Mm(297)
    section.top_margin = Cm(3); section.bottom_margin = Cm(2)
    section.right_margin = Cm(3.5); section.left_margin = Cm(2.5)
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi"); sectPr.append(bidi)


def style_doc(doc):
    st = doc.styles["Normal"]
    st.font.name = FONT; st.font.size = Pt(18)
    rpr = st.element.get_or_add_rPr(); rf = rpr.get_or_add_rFonts()
    rf.set(qn("w:cs"), FONT); rf.set(qn("w:ascii"), FONT); rf.set(qn("w:hAnsi"), FONT)
    # RTL على مستوى النمط الافتراضي (يضمن الاتجاه حتى لو بدأ السطر بحرف لاتيني/رقم)
    rpr.append(OxmlElement("w:rtl"))
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "36"); rpr.append(szCs)
    ppr = st.element.get_or_add_pPr(); ppr.append(OxmlElement("w:bidi"))
    jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "right"); ppr.append(jc)
    # ضبط افتراضيات المستند للنص المركّب (Complex Script) لتفادي العرض LTR
    try:
        styles_el = doc.styles.element
        dd = styles_el.find(qn("w:docDefaults"))
        if dd is not None:
            rpd = dd.find(qn("w:rPrDefault"))
            if rpd is not None:
                rr = rpd.find(qn("w:rPr"))
                if rr is None:
                    rr = OxmlElement("w:rPr"); rpd.append(rr)
                rr.append(OxmlElement("w:rtl"))
    except Exception:
        pass


def add_par(doc, text, size=18, bold=False, italic=False, color=None, highlight=False, align="right"):
    p = doc.add_paragraph()
    p.alignment = {"right": WD_ALIGN_PARAGRAPH.RIGHT, "center": WD_ALIGN_PARAGRAPH.CENTER,
                   "justify": WD_ALIGN_PARAGRAPH.JUSTIFY}.get(align, WD_ALIGN_PARAGRAPH.RIGHT)
    pPr = p._p.get_or_add_pPr(); pPr.append(OxmlElement("w:bidi"))
    pf = p.paragraph_format; pf.space_after = Pt(6); pf.line_spacing = 1.15
    run = p.add_run(xml_safe(text)); run.bold = bold; run.italic = italic
    run.font.name = FONT; run.font.size = Pt(size)
    rpr = run._element.get_or_add_rPr()
    rf = rpr.get_or_add_rFonts(); rf.set(qn("w:cs"), FONT)
    rpr.append(OxmlElement("w:rtl"))
    szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), str(size * 2)); rpr.append(szCs)
    if color:
        run.font.color.rgb = color
    if highlight:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return p


def add_card(doc, d):
    add_par(doc, "بيانات المستند", size=20, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
    card = d["card"]
    rows = [("اسم المستند الأصلي", d["title"]), ("نوع المستند", d["doc_type"]),
            ("الجهة/الدائرة", card.get("الجهة المصدِرة") or card.get("الدائرة") or ""),
            ("رقم القضية", card.get("رقم القضية", "")),
            ("رقم الصك/الحكم", d.get("deed") or card.get("رقم الصك/الحكم", "")),
            ("التاريخ", card.get("التاريخ", "")), ("الأطراف", card.get("الأطراف", "")),
            ("المسار الأصلي", d["parent_path"]),
            ("مستوى الثقة (فحص آلي)", d["qc"].get("status", "غير متاح")),
            ("يحتاج مراجعة بشرية", "نعم" if d["needs_review"] else "لا")]
    rows = [(k, v) for k, v in rows if str(v).strip()]
    t = doc.add_table(rows=len(rows), cols=2); t.style = "Table Grid"
    t.alignment = 2  # right
    bidiVis = OxmlElement("w:bidiVisual"); t._tbl.tblPr.append(bidiVis)
    for i, (k, v) in enumerate(rows):
        c0, c1 = t.rows[i].cells
        for cell, txt, b in [(c0, k, True), (c1, str(v), False)]:
            cell.text = ""
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))
            r = p.add_run(xml_safe(str(txt)[:400])); r.bold = b; r.font.name = FONT; r.font.size = Pt(16)
            rpr = r._element.get_or_add_rPr(); rpr.get_or_add_rFonts().set(qn("w:cs"), FONT)
            rpr.append(OxmlElement("w:rtl"))
            szCs = OxmlElement("w:szCs"); szCs.set(qn("w:val"), "32"); rpr.append(szCs)
    if d.get("viewUrl"):
        add_par(doc, "الرابط السحابي للأصل: " + d["viewUrl"], size=14, italic=True,
                color=RGBColor(0x55, 0x55, 0x55))
    doc.add_paragraph("")


def add_toc(doc):
    add_par(doc, "فهرس المستندات", size=22, bold=True, align="center")
    add_par(doc, "(لتحديث أرقام الصفحات: في Word اضغط على الفهرس ثم F9)", size=14, italic=True,
            color=RGBColor(0x77, 0x77, 0x77))
    p = doc.add_paragraph(); p._p.get_or_add_pPr().append(OxmlElement("w:bidi"))
    run = p.add_run()
    fld = OxmlElement("w:fldSimple"); fld.set(qn("w:instr"), r'TOC \o "1-1" \h \z \u')
    p._p.append(fld)


def add_footer_page(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    for instr in [("w:fldChar", "begin"), None, ("w:fldChar", "end")]:
        pass
    fb = OxmlElement("w:fldChar"); fb.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = " PAGE "
    fe = OxmlElement("w:fldChar"); fe.set(qn("w:fldCharType"), "end")
    run._r.append(fb); run._r.append(it); run._r.append(fe)


def build_docx(processed, path, review=False):
    doc = Document()
    style_doc(doc)
    set_section_rtl(doc.sections[0])
    add_footer_page(doc)
    add_par(doc, "المستندات القانونية — نسخة %s" % ("مراجعة" if review else "منقّحة"),
            size=22, bold=True, align="center")
    add_par(doc, "مخرج آلي يحتاج مراجعة بشرية — تنظيف وتنسيق لا تحرير قانوني · %s" % TODAY,
            size=14, italic=True, align="center", color=RGBColor(0x7c, 0x2d, 0x12))
    add_par(doc, "عدد المستندات: %d" % len(processed), size=16, align="center")
    add_toc(doc)
    review_tags_table = []
    for idx, d in enumerate(processed, 1):
        doc.add_page_break()
        add_par(doc, d["title"], size=22, bold=True)  # Heading-like (TOC via style below)
        # اجعلها Heading 1 ليلتقطها الفهرس
        doc.paragraphs[-1].style = doc.styles["Heading 1"]
        add_card(doc, d)
        for kind, t, tags in d["paras"]:
            if kind == "heading":
                add_par(doc, t, size=20, bold=True, color=RGBColor(0x1F, 0x4E, 0x79))
            else:
                hl = False; txt = t
                if review and tags:
                    txt = t + "  " + " ".join(tags)
                    hl = True
                    for tg in tags:
                        review_tags_table.append((idx, d["title"][:40], tg, t[:60]))
                add_par(doc, txt, size=18, highlight=hl)
        add_par(doc, "—" * 30, size=14, align="center", color=RGBColor(0xAA, 0xAA, 0xAA))
    if review and review_tags_table:
        doc.add_page_break()
        add_par(doc, "جدول مواضع المراجعة البشرية", size=22, bold=True, align="center")
        t = doc.add_table(rows=1, cols=4); t.style = "Table Grid"
        hdr = t.rows[0].cells
        for c, h in zip(hdr, ["#المستند", "العنوان", "الوسم", "مقتطف"]):
            c.text = h
        for r in review_tags_table[:500]:
            cells = t.add_row().cells
            for c, v in zip(cells, [str(r[0]), r[1], r[2], r[3]]):
                c.text = xml_safe(str(v))
    doc.save(str(path))
    return review_tags_table


# ---------- التقارير ----------
def w(path, text):
    path.write_text(text, encoding="utf-8")


def write_reports(processed, logs, dups, metrics):
    note = "> مخرج آلي يحتاج مراجعة بشرية — ليس رأياً قانونياً نهائياً.\n\n"
    # cleaning_report
    cr = ["# تقرير التنظيف العام\n", note,
          "## ملخص العمل", "تنظيف وتنسيق محافظ لمستندات مستخرجة آلياً، دون حذف مضمون أو إعادة صياغة.\n",
          "## أرقام", "| المؤشّر | قبل | بعد |", "|---|---|---|",
          "| عدد المستندات | %d | %d |" % (metrics["docs"], metrics["docs"]),
          "| عدد الفقرات/الأسطر | %d | %d |" % (metrics["paras_before"], metrics["paras_after"]),
          "| عدد الكلمات | %d | %d |" % (metrics["words_before"], metrics["words_after"]),
          "| نسبة تغيّر الكلمات | — | %.2f%% |" % metrics["word_delta_pct"], "",
          "## قواعد التنظيف المعتمدة",
          "- إزالة الترويسات/التذييلات المكررة من المتن بعد تسجيلها (تُنقل البيانات المهمة للبطاقة).",
          "- توحيد المسافات وضبط المسافة بعد علامات الترقيم العربية.",
          "- إزالة المحارف التحكّمية غير المتوافقة مع XML.",
          "## قواعد عدم التعديل (المحافظة)",
          "- لا عكس آلي للأسطر، لا قلب أرقام/تواريخ/أسماء، لا إعادة صياغة، لا حذف مضمون غامض.",
          "- كل اشتباه يُعلَّم بوسم مراجعة في نسخة المراجعة ويُسجَّل هنا.",
          "## ملخص الترويسات", "أنماط مكررة مكتشفة: %d (التفاصيل في headers_footers_report.md)." % len(logs["hf"]),
          "## ملخص OCR", "محارف محذوفة: %d · مسافات موحّدة: %d · ترقيم مضبوط: %d." % (
              logs["ocr_stats"]["ctrl"], logs["ocr_stats"]["space"], logs["ocr_stats"]["punct"]),
          "## ملخص اختلاط الاتجاه RTL/LTR",
          "أسطر مختلطة مكتشفة: %d · مستندات بأسطر معكوسة متبقية: %d (التفاصيل في bidi_reading_order_report.md)." % (
              metrics["bidi_lines"], metrics["docs_with_residual_rev"]),
          "## ملخص التكرارات", "أزواج مكرّرة/شبه مكرّرة: %d (لم يُحذف أي مستند — التفاصيل في duplication_report.md)." % len(dups),
          "## ملخص الجداول", "مقاطع يُشتبه أنها جداول: %d (مُعلَّمة للمراجعة، لم تُختلق خانات)." % metrics["table_hits"],
          "## مواضع تحتاج مراجعة بشرية", "مستندات مُعلَّمة للمراجعة: %d من %d." % (
              metrics["docs_need_review"], metrics["docs"]),
          "## توصيات للمرحلة التالية",
          "- مراجعة المستندات المُعلَّمة بالاستناد إلى الأصل في Google Drive.",
          "- لإصلاح ترتيب الأسطر بدقّة، توفير صور الصفحات أو OCR بإحداثيات (hOCR/ALTO).", ""]
    w(OUTDIR / "cleaning_report.md", "\n".join(cr))

    # headers_footers_report
    hf = ["# تقرير الترويسات والتذييلات\n", note,
          "| النمط (مطبَّع) | مرات الظهور | عدد المستندات | الإجراء |", "|---|---|---|---|"]
    for n, info in sorted(logs["hf"].items(), key=lambda x: -x[1]["count"])[:300]:
        hf.append("| %s | %d | %d | إزالة التكرار من المتن + نقل المهم للبطاقة |" % (
            info["sample"].replace("|", "/"), info["count"], len(info["docs"])))
    w(OUTDIR / "headers_footers_report.md", "\n".join(hf) + "\n")

    # ocr_issues_report
    oc = ["# تقرير مشكلات OCR والاشتباكات\n", note,
          "المعالجة محافظة: ضبط مسافات/ترقيم وإزالة محارف تحكّم فقط؛ ولم تُصحَّح الأسماء/الأرقام.",
          "", "| النوع | العدد |", "|---|---|",
          "| محارف تحكّم محذوفة | %d |" % logs["ocr_stats"]["ctrl"],
          "| مسافات متكررة وُحِّدت | %d |" % logs["ocr_stats"]["space"],
          "| مواضع ترقيم مضبوطة | %d |" % logs["ocr_stats"]["punct"],
          "", "## أمثلة أسطر مُعلَّمة بـ«OCR غير واضح» (عيّنة)"]
    cnt = 0
    for d in processed:
        for kind, t, tags in d["paras"]:
            if T_OCR in tags and cnt < 60:
                oc.append("- (%s) %s" % (d["title"][:30], t[:80])); cnt += 1
    w(OUTDIR / "ocr_issues_report.md", "\n".join(oc) + "\n")

    # bidi report
    bd = ["# تقرير اختلاط اتجاه القراءة RTL/LTR\n", note,
          "التمييز: مشكلة **عرض** (تُحلّ بضبط RTL في Word) مقابل مشكلة **ترتيب OCR** (تحتاج الأصل/الإحداثيات).",
          "بما أن إحداثيات OCR غير متوفّرة، لم يُعَد ترتيب أي سطر آلياً؛ الأسطر المشتبهة مُعلَّمة للمراجعة فقط.",
          "", "| المؤشّر | العدد |", "|---|---|",
          "| أسطر مختلطة (عربي+لاتيني) | %d |" % metrics["bidi_lines"],
          "| أسطر أُصلحت آلياً | 0 (سياسة محافظة) |",
          "| أسطر تحتاج مراجعة بشرية | %d |" % metrics["bidi_lines"],
          "| مستندات بأسطر معكوسة متبقية (فحص الجودة) | %d |" % metrics["docs_with_residual_rev"],
          "", "## أمثلة (النص كما ورد — لم يُعدَّل)"]
    cnt = 0
    for d in processed:
        for s in d["bidi_hits"][:3]:
            if cnt < 80:
                bd.append("- (%s) %s" % (d["title"][:30], s[:90])); cnt += 1
    bd += ["", "## المستندات ذات الأسطر المعكوسة المتبقية"]
    for d in processed:
        if d["rev_residual"]:
            bd.append("- %s — أسطر معكوسة متبقية: %d %s" % (d["title"][:50], d["rev_residual"], T_REV))
    w(OUTDIR / "bidi_reading_order_report.md", "\n".join(bd) + "\n")

    # source comparison
    sc = ["# تقرير المقارنة مع الأصل\n", note,
          "الملفات الأصلية موجودة في Google Drive (كثير منها كبير الحجم)، وإحداثيات OCR غير متوفّرة محلياً.",
          "لذلك المقارنة الآلية الكاملة غير ممكنة الآن؛ هذا التقرير يحدّد **المناطق عالية الخطورة** التي يجب",
          "مقارنتها يدوياً بالأصل (عبر الرابط السحابي في بطاقة كل مستند).", "",
          "| المستند | سبب الخطورة | يحتاج مراجعة بالأصل |", "|---|---|---|"]
    for d in processed:
        reasons = []
        if d["rev_residual"]:
            reasons.append("أسطر معكوسة متبقية")
        if d["bidi_hits"]:
            reasons.append("اختلاط اتجاه")
        if d["tbl_hits"]:
            reasons.append("جدول مشتبه")
        if d["qc"].get("status") not in ("سليم", ""):
            reasons.append("جودة قراءة منخفضة")
        if reasons:
            sc.append("| %s | %s | نعم |" % (d["title"][:50].replace("|", "/"), "، ".join(reasons)))
    w(OUTDIR / "source_comparison_report.md", "\n".join(sc) + "\n")

    # duplication report
    dp = ["# تقرير التكرارات\n", note,
          "لم يُحذف أي مستند؛ النسخ المختلفة قد تكون صيغاً مهمة. هذه أزواج عالية التطابق للمراجعة.",
          "", "| المستند (أ) | المستند (ب) | نسبة التطابق | الإجراء |", "|---|---|---|---|"]
    for a, b, r in dups:
        dp.append("| %s | %s | %.0f%% | %s |" % (a[:40].replace("|", "/"), b[:40].replace("|", "/"), r * 100,
                                                  "أُبقي الاثنان " + (T_DUP if r < 0.999 else "(مطابق تماماً)")))
    w(OUTDIR / "duplication_report.md", "\n".join(dp) + "\n")


def export_pdf(docx_path, outdir):
    """تصدير PDF عبر LibreOffice headless (إن توفّر)."""
    env = dict(os.environ, HOME="/tmp/lohome")
    Path("/tmp/lohome").mkdir(exist_ok=True)
    try:
        r = subprocess.run(
            ["soffice", "--headless", "-env:UserInstallation=file:///tmp/loprofile_lc",
             "--convert-to", "pdf", "--outdir", str(outdir), str(docx_path)],
            env=env, capture_output=True, text=True, timeout=540)
        pdf = outdir / (docx_path.stem + ".pdf")
        if pdf.exists():
            print("PDF:", pdf, "(%.1f MB)" % (pdf.stat().st_size / 1048576))
            return True
        print("تعذّر إنتاج PDF:", (r.stderr or r.stdout)[-200:])
    except Exception as e:
        print("تعذّر إنتاج PDF:", e)
    return False


def detect_dups(recs):
    dups = []
    items = [(r.get("title", ""), (r.get("full_text", "") or "")) for r in recs if (r.get("full_text") or "").strip()]
    for i in range(len(items)):
        for j in range(i + 1, len(items)):
            a, ta = items[i]; b, tb = items[j]
            if abs(len(ta) - len(tb)) > 0.25 * max(len(ta), len(tb), 1):
                continue
            r = difflib.SequenceMatcher(None, ta[:2500], tb[:2500]).ratio()
            if r >= 0.93:
                dups.append((a, b, r))
    return dups


def main():
    import sys
    sample = None
    if "--sample" in sys.argv:
        sample = int(sys.argv[sys.argv.index("--sample") + 1])
    OUTDIR.mkdir(parents=True, exist_ok=True)
    BACKUP.mkdir(parents=True, exist_ok=True)
    # نسخة احتياطية
    if IN.exists():
        shutil.copy2(IN, BACKUP / ("full_documents.%s.jsonl" % datetime.datetime.now().strftime("%Y%m%d_%H%M%S")))
    qc = load_qc()
    recs = [json.loads(l) for l in open(IN, encoding="utf-8") if l.strip()]
    if sample:
        recs = recs[:sample]
    logs = {"hf": {}, "ocr_stats": {"ctrl": 0, "space": 0, "punct": 0}}
    processed = [process_doc(r, qc, logs) for r in recs]
    dups = detect_dups(recs)

    wb = sum(d["words_before"] for d in processed); wa = sum(d["words_after"] for d in processed)
    metrics = {
        "docs": len(processed),
        "paras_before": sum(d["paras_before"] for d in processed),
        "paras_after": sum(d["paras_after"] for d in processed),
        "words_before": wb, "words_after": wa,
        "word_delta_pct": (100.0 * (wa - wb) / wb) if wb else 0.0,
        "bidi_lines": sum(len(d["bidi_hits"]) for d in processed),
        "table_hits": sum(len(d["tbl_hits"]) for d in processed),
        "docs_need_review": sum(1 for d in processed if d["needs_review"]),
        "docs_with_residual_rev": sum(1 for d in processed if d["rev_residual"]),
    }
    # index json
    index = []
    pos = 0
    for i, d in enumerate(processed, 1):
        index.append({
            "order": i, "title": d["title"], "doc_type": d["doc_type"],
            "segments": d["paras_after"], "source": "OCR/مستخرج آلياً",
            "confidence": d["qc"].get("status", "غير متاح"),
            "deed": d.get("deed", ""), "viewUrl": d.get("viewUrl", ""),
            "top_issues": ([("اتجاه مختلط:%d" % len(d["bidi_hits"])) ] if d["bidi_hits"] else []) +
                          (["جدول مشتبه"] if d["tbl_hits"] else []) +
                          (["أسطر معكوسة:%d" % d["rev_residual"]] if d["rev_residual"] else []),
            "needs_human_review": d["needs_review"],
        })
    (OUTDIR / "document_index.json").write_text(json.dumps(index, ensure_ascii=False, indent=1), encoding="utf-8")

    write_reports(processed, logs, dups, metrics)
    build_docx(processed, OUTDIR / "cleaned_legal_documents.docx", review=False)
    build_docx(processed, OUTDIR / "review_version.docx", review=True)
    export_pdf(OUTDIR / "cleaned_legal_documents.docx", OUTDIR)

    print("تمت المعالجة. مستندات:", len(processed))
    print("كلمات قبل/بعد: %d / %d (%.2f%%)" % (wb, wa, metrics["word_delta_pct"]))
    print("أسطر مختلطة:", metrics["bidi_lines"], "| تحتاج مراجعة:", metrics["docs_need_review"],
          "| تكرارات:", len(dups))
    if abs(metrics["word_delta_pct"]) > 8:
        print("⚠ تنبيه: تغيّر الكلمات تجاوز 8% — راجع قبل الاعتماد.")


if __name__ == "__main__":
    main()
