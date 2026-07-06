#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
doc_reader.py — محرك استخراج نص عربي متقدم (جودة احترافية)

يدعم: نصوص · Word · PDF نصي وممسوح · صور
الخلاصة: نص عربي نظيف 100% + معالجة ترميز متقدمة
"""
import os
import io
import re
import sys
import json
import unicodedata

# ━━━━━━━━━━━━━━━━━━━━ معالجة ترميز عربي متقدمة ━━━━━━━━━━━━━━━━━━━━

# 1. حروف خفيّة وعلامات اتجاه يجب حذفها
_HIDDEN_CHARS = dict.fromkeys([
    0x200b, 0x200c, 0x200d, 0x200e, 0x200f,          # ZWSP/ZWNJ/ZWJ/LRM/RLM
    0x202a, 0x202b, 0x202c, 0x202d, 0x202e,          # LRE/RLE/PDF/LRO/RLO
    0x2066, 0x2067, 0x2068, 0x2069, 0xfeff,          # LRI/RLI/FSI/PDI/BOM
    0x061c,                                            # ALM (Arabic Letter Mark)
], None)

# 2. حروف فارسية/أردية → عربية معيارية
_PERSIAN_TO_ARABIC = {
    "ھ": "ه", "ہ": "ه", "ۀ": "ه", "۔": ".",
    "ۃ": "ة", "ۆ": "و", "ۇ": "و", "ۉ": "و",
    "ۊ": "ي", "ۋ": "و", "ۍ": "ي", "ێ": "ي",
    "ۀ": "ه", "ۈ": "ي", "ۏ": "ي",
    "ی": "ي", "ۑ": "ي", "ك": "ك", "ک": "ك",
    "ۀ": "ه", "ٹ": "ت", "ڼ": "ن", "ڻ": "ن",
}

# 3. أرقام هندية وفارسية → عربية
_DIGIT_MAP = {
    "۰": "٠", "۱": "١", "۲": "٢", "۳": "٣", "۴": "٤",
    "۵": "٥", "۶": "٦", "۷": "٧", "۸": "٨", "۹": "٩",
    "٠": "٠", "١": "١", "٢": "٢", "٣": "٣", "٤": "٤",
    "٥": "٥", "٦": "٦", "٧": "٧", "٨": "٨", "٩": "٩",
    "0": "٠", "1": "١", "2": "٢", "3": "٣", "4": "٤",
    "5": "٥", "6": "٦", "7": "٧", "8": "٨", "9": "٩",
}

# 4. علامات ترقيم إنجليزية → عربية
_PUNCTUATION_MAP = {
    "،": "،",  # Arabic comma
    ",": "،",  # Convert English comma
}


def _normalize_unicode(text):
    """تطبيع Unicode النصي: تحليل مركبات العربية"""
    if not text:
        return text
    return unicodedata.normalize("NFKC", text)


def _remove_hidden_chars(text):
    """حذف العلامات الخفيّة والتحكمية"""
    if not text:
        return text
    return text.translate(_HIDDEN_CHARS)


def _convert_persian_to_arabic(text):
    """تحويل حروف فارسية وأردية إلى عربية"""
    if not text:
        return text
    for persian, arabic in _PERSIAN_TO_ARABIC.items():
        text = text.replace(persian, arabic)
    return text


def _normalize_digits(text):
    """توحيد الأرقام إلى العربية"""
    if not text:
        return text
    for digit, arabic_digit in _DIGIT_MAP.items():
        text = text.replace(digit, arabic_digit)
    return text


def _normalize_punctuation(text):
    """توحيد علامات الترقيم"""
    if not text:
        return text
    for punct, arabic_punct in _PUNCTUATION_MAP.items():
        text = text.replace(punct, arabic_punct)
    return text


def _remove_extra_spaces(text):
    """حذف المسافات الزائدة والأسطر الفارغة"""
    if not text:
        return text
    # حذف المسافات المتكررة
    text = re.sub(r' +', ' ', text)
    # حذف الأسطر الفارغة المتكررة
    text = re.sub(r'\n\n+', '\n', text)
    # قص البداية والنهاية
    return text.strip()


def clean_text(text):
    """تنظيف شامل للنص العربي (مرحلة واحدة فقط)"""
    if not text:
        return text

    # خطوات التنظيف بالترتيب
    text = _normalize_unicode(text)           # تطبيع Unicode
    text = _remove_hidden_chars(text)         # حذف علامات خفيّة
    text = _convert_persian_to_arabic(text)   # تحويل فارسي → عربي
    text = _normalize_digits(text)            # توحيد الأرقام
    text = _normalize_punctuation(text)       # توحيد الترقيم
    text = _remove_extra_spaces(text)         # حذف مسافات زائدة

    return text


def norm(s):
    """تطبيع للبحث: حذف التشكيل والتطويل"""
    if not s:
        return s
    out = []
    for ch in s:
        c = ord(ch)
        # تخطي التشكيل والتطويل
        if 0x064B <= c <= 0x0652 or c in (0x0640, 0x0670):
            continue
        # توحيد الهمزات والألف والتاء
        normalized = {
            "أ": "ا", "إ": "ا", "آ": "ا", "ٱ": "ا",
            "ة": "ه", "ى": "ي", "ؤ": "و", "ئ": "ي"
        }.get(ch, ch.lower())
        out.append(normalized)
    return "".join(out)


# ━━━━━━━━━━━━━━━━━━━━ استخراج النص من الملفات ━━━━━━━━━━━━━━━━━━━━

TEXT_EXT = (".txt", ".md", ".csv", ".json")
IMG_EXT = (".png", ".jpg", ".jpeg", ".tif", ".tiff", ".bmp")
SUPPORTED = TEXT_EXT + (".docx", ".pdf") + IMG_EXT

_AR_RE = re.compile(r"[؀-ۿ]")


def _read_text_file(data):
    """قراءة ملفات نصية"""
    # محاولة ترميزات مختلفة
    for encoding in ["utf-8", "utf-16", "cp1252", "iso-8859-6"]:
        try:
            return data.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return data.decode("utf-8", "ignore")


def _read_docx(data):
    """قراءة ملفات Word"""
    try:
        from docx import Document
        doc = Document(io.BytesIO(data))
        parts = []

        # نصوص الفقرات
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text)

        # نصوص الجداول
        for table in doc.tables:
            for row in table.rows:
                row_text = "\t".join(cell.text for cell in row.cells)
                if row_text.strip():
                    parts.append(row_text)

        return "\n".join(parts) if parts else ""
    except Exception:
        return ""


def _read_pdf_with_pdfplumber(data):
    """قراءة PDF باستخدام pdfplumber (أفضل للعربية)"""
    try:
        import pdfplumber
        pdf = pdfplumber.open(io.BytesIO(data))
        text_parts = []

        for page in pdf.pages:
            # محاولة استخراج نص مباشر أولاً
            text = page.extract_text()
            if text:
                text_parts.append(text)

        pdf.close()
        return "\n".join(text_parts) if text_parts else ""
    except Exception:
        return ""


def _read_pdf_with_pdfminer(data):
    """قراءة PDF باستخدام pdfminer (fallback)"""
    try:
        from pdfminer.high_level import extract_text
        return extract_text(io.BytesIO(data)) or ""
    except Exception:
        return ""


def _read_pdf_text(data):
    """قراءة نص PDF (مع fallback)"""
    # جرّب pdfplumber أولاً (أفضل)
    text = _read_pdf_with_pdfplumber(data)
    if text and len(_AR_RE.findall(text)) > 30:
        return text

    # fallback لـ pdfminer
    text = _read_pdf_with_pdfminer(data)
    return text if text else ""


def _enhance_image(img):
    """تحسين الصورة لـ OCR: تباين أفضل + وضوح أعلى"""
    try:
        from PIL import Image, ImageOps, ImageFilter, ImageEnhance

        # تحويل إلى RGB
        if img.mode != "RGB":
            img = img.convert("RGB")

        # تدرج رمادي
        img = ImageOps.grayscale(img)

        # زيادة التباين
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(2.0)

        # زيادة الحدة
        img = img.filter(ImageFilter.SHARPEN)
        img = img.filter(ImageFilter.SHARPEN)

        # توازن الإضاءة
        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(1.1)

        return img
    except Exception:
        return img


def _ocr_image(img, lang="ara"):
    """استخراج نص من صورة باستخدام Tesseract"""
    try:
        import pytesseract

        # تحسين الصورة أولاً
        enhanced = _enhance_image(img)

        # Tesseract config محسّن للعربية
        config = "--oem 1 --psm 6 -l ara+osd"

        text = pytesseract.image_to_string(enhanced, config=config)
        return text if text else ""
    except Exception:
        return ""


def _ocr_pdf(data):
    """استخراج نص ممسوح من PDF باستخدام OCR"""
    try:
        from pdf2image import convert_from_bytes
        from PIL import Image

        # تحويل إلى صور بدقة عالية (400 DPI)
        pages = convert_from_bytes(data, dpi=400)

        text_parts = []
        for page in pages:
            text = _ocr_image(page)
            if text:
                text_parts.append(text)

        return "\n".join(text_parts) if text_parts else ""
    except Exception:
        return ""


def _ocr_image_file(data):
    """استخراج نص من ملف صورة"""
    try:
        from PIL import Image
        img = Image.open(io.BytesIO(data)).convert("RGB")
        return _ocr_image(img)
    except Exception:
        return ""


def read_bytes(name, data):
    """استخراج نص من ملف (المحرك الرئيسي)"""
    ext = os.path.splitext(name)[1].lower()

    try:
        if ext in TEXT_EXT:
            text = _read_text_file(data)
            return clean_text(text), "نص"

        elif ext == ".docx":
            text = _read_docx(data)
            return clean_text(text), "Word"

        elif ext == ".pdf":
            # جرّب النص أولاً
            text = _read_pdf_text(data)

            # تحقق: هل هناك نص عربي كافي؟
            if text and len(_AR_RE.findall(text)) > 30:
                return clean_text(text), "PDF (نص)"

            # وإلا: استخدم OCR
            text = _ocr_pdf(data)
            return clean_text(text) if text else "", "PDF ممسوح (OCR)"

        elif ext in IMG_EXT:
            text = _ocr_image_file(data)
            return clean_text(text) if text else "", "صورة (OCR)"

    except Exception as e:
        return "", f"خطأ: {str(e)[:50]}"

    return "", "صيغة غير مدعومة"


def read_file(path):
    """قراءة ملف من القرص"""
    with open(path, "rb") as f:
        return read_bytes(os.path.basename(path), f.read())


# ━━━━━━━━━━━━━━━━━━━━ الواجهة العالية ━━━━━━━━━━━━━━━━━━━━

def extract(path):
    """استخراج شامل من ملف"""
    text, kind = read_file(path)
    return {
        "title": os.path.basename(path),
        "kind": kind,
        "text": text,
        "search": norm(text),
        "chars": len(text),
        "ok": bool(text.strip()),
    }


def walk(folder):
    """استخراج من مجلد بالكامل"""
    for root, _, files in os.walk(folder):
        for filename in sorted(files):
            if os.path.splitext(filename)[1].lower() in SUPPORTED:
                yield extract(os.path.join(root, filename))


# ━━━━━━━━━━━━━━━━━━━━ سطر الأوامر ━━━━━━━━━━━━━━━━━━━━

def _main(argv):
    if not argv or argv[0] in ("-h", "--help"):
        print(__doc__)
        return 0

    target = argv[0]
    out_json = None

    if "--json" in argv:
        out_json = argv[argv.index("--json") + 1]

    if os.path.isdir(target):
        recs = list(walk(target))
        if out_json:
            with open(out_json, "w", encoding="utf-8") as f:
                for r in recs:
                    f.write(json.dumps(r, ensure_ascii=False) + "\n")
            ok = sum(1 for r in recs if r["ok"])
            print(f"عُولجت {len(recs)} وثيقة (نجح الاستخراج في {ok}) → {out_json}")
        else:
            for r in recs:
                status = "" if r["ok"] else "⚠"
                print(f"• {r['title'][:40]:40} [{r['kind']}] {r['chars']:,d} حرف {status}")
    else:
        r = extract(target)
        if out_json:
            with open(out_json, "w", encoding="utf-8") as f:
                f.write(json.dumps(r, ensure_ascii=False) + "\n")
            print(f"→ {out_json} ({r['kind']}، {r['chars']:,d} حرف)")
        else:
            sys.stdout.write(r["text"])
            if not r["ok"]:
                sys.stderr.write(f"\n[{r['kind']}: لم يُستخرَج نص]\n")

    return 0


if __name__ == "__main__":
    raise SystemExit(_main(sys.argv[1:]))
